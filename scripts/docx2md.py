#!/usr/bin/env python3
"""Convert Word (.docx) documents to Markdown.

Features:
  - Heading styles -> Markdown headings (# ## ### ...)
  - Tables -> HTML tables with rowspan/colspan for merged cells
  - Ordered/unordered lists -> Markdown lists (nested supported)
  - Inline formatting: bold, italic, strikethrough, code, hyperlinks
  - Chinese numbering detection (附件, 一二三, （一）（二）)
  - Preserves document element order (paragraphs, tables, lists)

Usage:
  python docx2md.py input.docx [-o output.md]

Can also be imported as a module:
  from docx2md import convert
  convert("input.docx", "output.md")

Requires: python-docx (pip install python-docx)
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------

def _escape_md(text: str) -> str:
    """Escape Markdown special characters in plain text."""
    if not text:
        return ""
    for ch in (r"\\", "`", "*", "_", "[", "]", "#"):
        text = text.replace(ch, "\\" + ch)
    return text


def _para_inlines(para) -> str:
    """Build inline Markdown from a paragraph's runs + hyperlinks."""
    parts = []
    for child in list(para._element):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "r":
            rPr = child.find(qn("w:rPr"))
            is_bold = is_italic = is_strike = is_code = False
            if rPr is not None:
                b_el = rPr.find(qn("w:b"))
                if b_el is not None and b_el.get(qn("w:val")) != "0":
                    is_bold = True
                i_el = rPr.find(qn("w:i"))
                if i_el is not None and i_el.get(qn("w:val")) != "0":
                    is_italic = True
                strike_el = rPr.find(qn("w:strike"))
                if strike_el is not None and strike_el.get(qn("w:val")) != "0":
                    is_strike = True
                rfonts = rPr.find(qn("w:rFonts"))
                if rfonts is not None:
                    af = (rfonts.get(qn("w:ascii")) or "").lower()
                    if af in ("consolas", "courier", "courier new", "monospace",
                              "source code pro", "fira code", "menlo"):
                        is_code = True

            t_el = child.find(qn("w:t"))
            text = t_el.text if t_el is not None and t_el.text else ""
            if not text:
                continue

            if is_code:
                parts.append(f"`{text}`")
            else:
                result = _escape_md(text)
                if result.strip():
                    if is_strike:
                        result = f"~~{result}~~"
                    if is_italic:
                        result = f"*{result}*"
                    if is_bold:
                        result = f"**{result}**"
                parts.append(result)

        elif tag == "hyperlink":
            r_id = child.get(qn("r:id"))
            if r_id and hasattr(para.part, "rels"):
                try:
                    rel = para.part.rels[r_id]
                    url = rel.target_ref
                    link_text = ""
                    for r_el in child.findall(qn("w:r")):
                        t_el = r_el.find(qn("w:t"))
                        if t_el is not None and t_el.text:
                            link_text += t_el.text
                    if link_text and url:
                        parts.append(f"[{link_text}]({url})")
                except Exception:
                    pass

    return "".join(parts)

# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

# Chinese numbering patterns for heading detection
_RE_FJ = re.compile(r"^附件\s*[一二三四五六七八九十\d]+")
_RE_CN_L1 = re.compile(r"^([一二三四五六七八九十]+)[、．.]\s*(.*)")
_RE_CN_L2_FULL = re.compile(r"^（([一二三四五六七八九十]+)）\s*(.*)")
_RE_CN_L2_HALF = re.compile(r"^\(([一二三四五六七八九十]+)\)\s*(.*)")
_RE_ARABIC_DOT = re.compile(r"^(\d+)[.．、]\s*(.*)")
_RE_ARABIC_PAREN = re.compile(r"^\((\d+)\)\s*(.*)")


def _para_font_size(para):
    """Return the font size (in pt) of the first non-empty run, or None."""
    for r in para.runs:
        if r.text.strip() and r.font and r.font.size:
            return r.font.size.pt
    return None


def _all_runs_bold(para):
    """Return True if every non-empty run is bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _heading_level(para) -> int:
    """Return heading level (1-9) or 0 if not a heading.

    Detection order:
      1. Explicit Word heading style (Title / Subtitle / Heading N)
      2. outlineLevel in style or paragraph XML
      3. Chinese appendix marker  附件1 / 附件一  -> H1
      4. Document title: centered + bold + large font (>=20 pt) -> H1
      5. Chinese L1 numbering  一、二、三...  -> H2
      6. Chinese L2 numbering  （一）（二）...  -> H3  (any style)
      7. Chinese L2 numbering  (一)(二)...  -> H3  (only non-Body-Text)
      8. Font-size heuristic: short + bold + >=15 pt in Normal -> H3
    """
    style = para.style
    style_name = (style.name or "").lower() if style else ""
    text = para.text.strip()
    if not text:
        return 0

    # ---- 1. Explicit heading style ----
    if style_name == "title":
        return 1
    if style_name == "subtitle":
        return 2
    if style_name.startswith("heading"):
        try:
            n = int(style_name.replace("heading", "").strip())
            return min(n, 9)
        except ValueError:
            pass

    # ---- 2. outlineLevel (style -> paragraph) ----
    if style and hasattr(style, "element"):
        pPr = style.element.find(qn("w:pPr"))
        if pPr is not None:
            outlineLvl = pPr.find(qn("w:outlineLevel"))
            if outlineLvl is not None:
                val = outlineLvl.get(qn("w:val"))
                if val is not None:
                    return int(val) + 1

    pPr_el = para._element.find(qn("w:pPr"))
    if pPr_el is not None:
        outlineLvl = pPr_el.find(qn("w:outlineLevel"))
        if outlineLvl is not None:
            val = outlineLvl.get(qn("w:val"))
            if val is not None:
                return int(val) + 1

    is_body_text = style_name in ("body text",)
    is_normal = style_name in ("normal", "")

    # ---- 3. 附件N -> H1 ----
    if _RE_FJ.match(text) and len(text) < 15:
        return 1

    # ---- 4. Document title: centered + bold + large font -> H1 ----
    align = para.alignment
    is_centered = align is not None and str(align) in ("CENTER (1)", "1")
    font_size = _para_font_size(para)
    if is_centered and _all_runs_bold(para) and font_size and font_size >= 20:
        return 1

    # ---- 5. Chinese L1:  一、二、三、... -> H2 ----
    if _RE_CN_L1.match(text):
        return 2

    # ---- 6. Chinese L2 full-width:  （一）（二）... -> H3 (any style) ----
    if _RE_CN_L2_FULL.match(text):
        return 3

    # ---- 7. Chinese L2 half-width:  (一)(二)... -> H3 only if not Body Text ----
    if not is_body_text and _RE_CN_L2_HALF.match(text):
        return 3

    # ---- 8. Font-size heuristic: short bold paragraph in Normal -> H3 ----
    if is_normal and _all_runs_bold(para) and font_size and font_size >= 15:
        if len(text) < 40:
            return 3

    return 0


# ---------------------------------------------------------------------------
# Body-text list pattern detection
# ---------------------------------------------------------------------------

def _body_text_list_pattern(para) -> str | None:
    """If a Body-Text paragraph starts with a numbering pattern that looks
    like a sub-heading / list item, return the marker text to use.

    Returns None if the paragraph should stay as a normal paragraph.
    """
    style_name = (para.style.name or "").lower() if para.style else ""
    if style_name not in ("body text",):
        return None

    text = para.text.strip()
    if not text:
        return None

    # (一) (二) (三) ...  (half-width parens, Chinese numerals)
    if _RE_CN_L2_HALF.match(text):
        return "-"

    # 1. 2. 3. ... (Arabic numeral + dot)
    if _RE_ARABIC_DOT.match(text):
        return "-"

    # (1) (2) (3) ... (Arabic numeral in parens)
    if _RE_ARABIC_PAREN.match(text):
        return "-"

    return None

# ---------------------------------------------------------------------------
# List detection  (checks both paragraph-level and style-level numPr)
# ---------------------------------------------------------------------------

def _get_num_info(doc, para) -> dict | None:
    """Extract numbering (list) info from a paragraph or its style.

    Returns dict with keys: numId, ilvl, is_ordered  or None.
    """
    numId = None
    ilvl = 0

    # 1) Check paragraph-level numPr
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            numId_el = numPr.find(qn("w:numId"))
            ilvl_el = numPr.find(qn("w:ilvl"))
            if numId_el is not None:
                numId = int(numId_el.get(qn("w:val"), "0"))
                ilvl = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0

    # 2) Fallback: check style definition chain
    if numId is None and para.style is not None:
        style_el = para.style.element
        if style_el is not None:
            spPr = style_el.find(qn("w:pPr"))
            if spPr is not None:
                numPr = spPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId_el = numPr.find(qn("w:numId"))
                    ilvl_el = numPr.find(qn("w:ilvl"))
                    if numId_el is not None:
                        numId = int(numId_el.get(qn("w:val"), "0"))
                        ilvl = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0

    if numId is None or numId == 0:
        return None

    # Determine ordered vs unordered from numbering definitions
    is_ordered = False
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is not None:
            numbering_xml = numbering_part._element
            for num_el in numbering_xml.findall(qn("w:num")):
                nid = num_el.get(qn("w:numId"))
                if nid == str(numId):
                    abs_ref = num_el.find(qn("w:abstractNumId"))
                    if abs_ref is not None:
                        abs_id = abs_ref.get(qn("w:val"))
                        for abs_el in numbering_xml.findall(qn("w:abstractNum")):
                            if abs_el.get(qn("w:abstractNumId")) == abs_id:
                                for lvl in abs_el.findall(qn("w:lvl")):
                                    if lvl.get(qn("w:ilvl")) == str(ilvl):
                                        numFmt = lvl.find(qn("w:numFmt"))
                                        if numFmt is not None:
                                            fmt = numFmt.get(qn("w:val"), "")
                                            is_ordered = fmt == "decimal"
                                        break
                        break
        # Fallback: style name heuristic
        if not is_ordered:
            sn = (para.style.name or "").lower()
            if "number" in sn or "ordered" in sn:
                is_ordered = True
    except Exception:
        sn = (para.style.name or "").lower()
        if "number" in sn or "ordered" in sn:
            is_ordered = True

    return {"numId": numId, "ilvl": ilvl, "is_ordered": is_ordered}

# ---------------------------------------------------------------------------
# Table -> HTML  (merged-cell support via _tc identity + vMerge/gridSpan)
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    """Get text content of a table cell (flattened)."""
    lines = []
    for p in cell.paragraphs:
        t = _para_inlines(p)
        if t.strip():
            lines.append(t.strip())
    return "<br>".join(lines) if lines else ""


def _table_to_html(table) -> str:
    """Convert a Word table to an HTML table with rowspan/colspan."""
    rows_data = []
    for row in table.rows:
        cells = [cell for cell in row.cells]
        rows_data.append(cells)

    n_rows = len(rows_data)
    if n_rows == 0:
        return ""
    n_cols = max(len(r) for r in rows_data) if rows_data else 0
    if n_cols == 0:
        return ""

    # Track which (row, col) positions are already covered by a span
    covered = set()
    # Track which _tc objects we have already emitted (by id)
    emitted_tc_ids = set()

    html = ["<table>", "  <tbody>"]

    for r, row_cells in enumerate(rows_data):
        html.append("    <tr>")
        for c in range(n_cols):
            if (r, c) in covered:
                continue
            if c >= len(row_cells):
                break

            cell = row_cells[c]
            tc = cell._tc
            tc_id = id(tc)

            # If this _tc was already emitted in a previous row,
            # it's a vertical merge continuation -> skip
            if tc_id in emitted_tc_ids:
                covered.add((r, c))
                continue

            emitted_tc_ids.add(tc_id)

            # --- colspan from gridSpan attribute ---
            tcPr = tc.find(qn("w:tcPr"))
            colspan = 1
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    colspan = int(gs.get(qn("w:val"), "1"))

            # --- rowspan: count rows where same _tc appears ---
            rowspan = 1
            for nr in range(r + 1, n_rows):
                if c < len(rows_data[nr]) and id(rows_data[nr][c]._tc) == tc_id:
                    rowspan += 1
                    covered.add((nr, c))
                else:
                    break

            # Mark horizontally covered positions
            for dc in range(1, colspan):
                if c + dc < n_cols:
                    covered.add((r, c + dc))

            # Build cell HTML
            text = _cell_text(cell)
            attrs = ""
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'

            tag = "th" if r == 0 else "td"
            html.append(f"      <{tag}{attrs}>{text}</{tag}>")

        html.append("    </tr>")

    html.append("  </tbody>")
    html.append("</table>")
    return "\n".join(html)

# ---------------------------------------------------------------------------
# Main document body walker (preserves element order)
# ---------------------------------------------------------------------------

def convert_body(doc) -> str:
    """Walk the document body in order, converting each element."""
    body = doc.element.body
    lines: list[str] = []
    list_counters: dict[int, int] = {}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            # --- Heading ---
            hlevel = _heading_level(para)
            if hlevel > 0:
                text = _para_inlines(para)
                if text.strip():
                    lines.append("")
                    lines.append("#" * hlevel + " " + text)
                    lines.append("")
                    list_counters.clear()
                continue

            # --- List item ---
            num_info = _get_num_info(doc, para)
            if num_info is not None:
                indent = "  " * num_info["ilvl"]
                if num_info["is_ordered"]:
                    key = num_info["numId"]
                    list_counters[key] = list_counters.get(key, 0) + 1
                    marker = f"{list_counters[key]}."
                else:
                    marker = "-"
                text = _para_inlines(para)
                lines.append(f"{indent}{marker} {text}")
                continue
            else:
                list_counters.clear()

            # --- Body-text numbering pattern -> Markdown list item ---
            bt_marker = _body_text_list_pattern(para)
            if bt_marker is not None:
                text = _para_inlines(para)
                lines.append(f"{bt_marker} {text}")
                continue

            # --- Normal paragraph ---
            text = _para_inlines(para)
            if text.strip():
                lines.append("")
                lines.append(text)
                lines.append("")

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is child:
                    lines.append("")
                    lines.append(_table_to_html(table))
                    lines.append("")
                    list_counters.clear()
                    break

    return "\n".join(lines)

# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def convert(docx_path: str, output_path: str | None = None) -> str:
    """Convert a .docx file to Markdown and return the Markdown string."""
    doc = Document(docx_path)
    md = convert_body(doc)

    # Clean up excessive blank lines
    while "\n\n\n" in md:
        md = md.replace("\n\n\n", "\n\n")
    md = md.strip() + "\n"

    if output_path:
        Path(output_path).write_text(md, encoding="utf-8")
        print(f"OK: {docx_path} -> {output_path}", file=sys.stderr)
    return md


def main():
    ap = argparse.ArgumentParser(
        description="Convert Word (.docx) to Markdown",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("input", help="Input .docx file path")
    ap.add_argument("-o", "--output", help="Output .md file (default: <input>.md)")
    args = ap.parse_args()

    inp = Path(args.input)
    if not inp.exists():
        print(f"Error: file not found: {inp}", file=sys.stderr)
        sys.exit(1)

    out = args.output or str(inp.with_suffix(".md"))
    convert(str(inp), out)


if __name__ == "__main__":
    main()
